# import streamlit as st
# import tempfile
# import unicodedata
# from docx import Document
# from PyPDF2 import PdfReader
# from fpdf import FPDF
# from langdetect import detect
# from datetime import datetime
# from googletrans import Translator
# import firebase_admin
# from firebase_admin import credentials, firestore
# from deep_translator import GoogleTranslator

# # Initialize Firebase
# if not firebase_admin._apps:
#     cred = credentials.Certificate("serviceAccountKey.json")
#     firebase_admin.initialize_app(cred)
# db = firestore.client()

# translator = Translator()

# # Language map
# LANG_MAP = {
#     'Afrikaans': 'af', 'Albanian': 'sq', 'Amharic': 'am', 'Arabic': 'ar',
#     'Armenian': 'hy', 'Assamese': 'as', 'Aymara': 'ay', 'Azerbaijani': 'az',
#     'Bambara': 'bm', 'Basque': 'eu', 'Belarusian': 'be', 'Bengali': 'bn',
#     'Bhojpuri': 'bho', 'Bosnian': 'bs', 'Bulgarian': 'bg', 'Catalan': 'ca',
#     'Cebuano': 'ceb', 'Chichewa': 'ny', 'Chinese (Simplified)': 'zh-CN', 'Chinese (Traditional)': 'zh-TW',
#     'Corsican': 'co', 'Croatian': 'hr', 'Czech': 'cs', 'Danish': 'da',
#     'Dhivehi': 'dv', 'Dogri': 'doi', 'Dutch': 'nl', 'English': 'en',
#     'Esperanto': 'eo', 'Estonian': 'et', 'Ewe': 'ee', 'Filipino': 'tl',
#     'Finnish': 'fi', 'French': 'fr', 'Frisian': 'fy', 'Galician': 'gl',
#     'Georgian': 'ka', 'German': 'de', 'Greek': 'el', 'Guarani': 'gn',
#     'Gujarati': 'gu', 'Haitian Creole': 'ht', 'Hausa': 'ha', 'Hawaiian': 'haw',
#     'Hebrew': 'iw', 'Hindi': 'hi', 'Hmong': 'hmn', 'Hungarian': 'hu',
#     'Icelandic': 'is', 'Igbo': 'ig', 'Ilocano': 'ilo', 'Indonesian': 'id',
#     'Irish': 'ga', 'Italian': 'it', 'Japanese': 'ja', 'Javanese': 'jw',
#     'Kannada': 'kn', 'Kazakh': 'kk', 'Khmer': 'km', 'Kinyarwanda': 'rw',
#     'Konkani': 'gom', 'Korean': 'ko', 'Krio': 'kri', 'Kurdish (Kurmanji)': 'ku',
#     'Kurdish (Sorani)': 'ckb', 'Kyrgyz': 'ky', 'Lao': 'lo', 'Latin': 'la',
#     'Latvian': 'lv', 'Lingala': 'ln', 'Lithuanian': 'lt', 'Luganda': 'lg',
#     'Luxembourgish': 'lb', 'Macedonian': 'mk', 'Maithili': 'mai', 'Malagasy': 'mg',
#     'Malay': 'ms', 'Malayalam': 'ml', 'Maltese': 'mt', 'Maori': 'mi',
#     'Marathi': 'mr', 'Meiteilon (Manipuri)': 'mni-Mtei', 'Mizo': 'lus', 'Mongolian': 'mn',
#     'Myanmar': 'my', 'Nepali': 'ne', 'Norwegian': 'no', 'Odia (Oriya)': 'or',
#     'Oromo': 'om', 'Pashto': 'ps', 'Persian': 'fa', 'Polish': 'pl',
#     'Portuguese': 'pt', 'Punjabi': 'pa', 'Quechua': 'qu', 'Romanian': 'ro',
#     'Russian': 'ru', 'Samoan': 'sm', 'Sanskrit': 'sa', 'Scots Gaelic': 'gd',
#     'Sepedi': 'nso', 'Serbian': 'sr', 'Sesotho': 'st', 'Shona': 'sn',
#     'Sindhi': 'sd', 'Sinhala': 'si', 'Slovak': 'sk', 'Slovenian': 'sl',
#     'Somali': 'so', 'Spanish': 'es', 'Sundanese': 'su', 'Swahili': 'sw',
#     'Swedish': 'sv', 'Tajik': 'tg', 'Tamil': 'ta', 'Tatar': 'tt',
#     'Telugu': 'te', 'Thai': 'th', 'Tigrinya': 'ti', 'Tsonga': 'ts',
#     'Turkish': 'tr', 'Turkmen': 'tk', 'Twi': 'ak', 'Ukrainian': 'uk',
#     'Urdu': 'ur', 'Uyghur': 'ug', 'Uzbek': 'uz', 'Vietnamese': 'vi',
#     'Welsh': 'cy', 'Xhosa': 'xh', 'Yiddish': 'yi', 'Yoruba': 'yo',
#     'Zulu': 'zu'
# }

# RTL_LANGS = ['ar', 'iw', 'fa', 'ur']

# def get_language_code(name):
#     return LANG_MAP.get(name, "en")

# def detect_language(text):
#     return detect(text)

# def translate_text(text, src, dest):
#     return GoogleTranslator(source=src, target=dest).translate(text)

# def extract_text_from_pdf(file):
#     reader = PdfReader(file)
#     return "\n".join([page.extract_text() or '' for page in reader.pages])

# def extract_text_from_docx(file):
#     doc = Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# def normalize_text(text):
#     return unicodedata.normalize("NFKC", text)

# def create_translated_pdf(text, output_path, lang_code='en'):
#     pdf = FPDF(orientation="R" if lang_code in RTL_LANGS else "P")
#     pdf.add_page()
#     pdf.set_auto_page_break(auto=True, margin=15)

#     # Load Unicode font
#     font_path = "assets/fonts/DejaVuSans.ttf"
#     pdf.add_font("DejaVu", "", font_path, uni=True)
#     pdf.set_font("DejaVu", size=12)

#     lines = text.split("\n")
#     for line in lines:
#         line = normalize_text(line)
#         if lang_code in RTL_LANGS:
#             line = line[::-1]  # crude RTL support (mirroring)
#         pdf.multi_cell(0, 10, line)

#     pdf.output(output_path)

# def create_translated_docx(text, output_path):
#     doc = Document()
#     for para in text.split('\n'):
#         doc.add_paragraph(normalize_text(para))
#     doc.save(output_path)

# def save_txt(text, output_path):
#     with open(output_path, 'w', encoding='utf-8') as f:
#         f.write(normalize_text(text))

# def run(uid=None):
#     st.title("📄 Document Translator")
#     st.write("Upload a **PDF**, **DOCX**, **TXT**, **MD**, or **RTF** file to translate its contents.")

#     uploaded_file = st.file_uploader("Choose a document", type=["pdf", "docx", "txt", "md", "rtf"])
#     target_language = st.selectbox("Translate to", options=list(LANG_MAP.keys()))

#     if uploaded_file and target_language:
#         with st.spinner("🔍 Reading document..."):
#             ext = uploaded_file.name.split('.')[-1].lower()
#             if ext == "pdf":
#                 text = extract_text_from_pdf(uploaded_file)
#             elif ext == "docx":
#                 text = extract_text_from_docx(uploaded_file)
#             elif ext in ["txt", "md", "rtf"]:
#                 text = uploaded_file.read().decode("utf-8")
#             else:
#                 st.error("❌ Unsupported file format.")
#                 return

#         st.success("✅ Text extracted successfully!")
#         st.text_area("Extracted Text Preview", text[:2000], height=200)

#         if st.button("🌍 Translate Document"):
#             with st.spinner("🌐 Translating..."):
#                 src_lang = detect_language(text)
#                 tgt_lang_code = get_language_code(target_language)
#                 translated = translate_text(text, src=src_lang, dest=tgt_lang_code)

#                 with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as out_file:
#                     out_path = out_file.name

#                 if ext == "pdf":
#                     create_translated_pdf(translated, out_path, lang_code=tgt_lang_code)
#                 elif ext == "docx":
#                     create_translated_docx(translated, out_path)
#                 else:
#                     save_txt(translated, out_path)

#                 # Firestore save
#                 user_email = st.session_state.get("user_email", "anonymous")
#                 db.collection("translation_history").add({
#                     "user": user_email,
#                     "filename": uploaded_file.name,
#                     "filetype": ext,
#                     "source_lang": src_lang,
#                     "target_lang": tgt_lang_code,
#                     "timestamp": datetime.utcnow(),
#                     "text_preview": text[:200]
#                 })

#                 st.success("✅ Document translated and saved!")
#                 with open(out_path, "rb") as f:
#                     st.download_button(
#                         f"📥 Download Translated {ext.upper()}",
#                         f,
#                         file_name=f"translated_{uploaded_file.name}"
#                     )

#     if st.button("⬅️ Back to Dashboard"):
#         st.session_state["current_page"] = "🏠 Dashboard"
#         st.rerun()

import streamlit as st
import os
import tempfile
from docx import Document
from PyPDF2 import PdfReader
from fpdf import FPDF
from langdetect import detect
from googletrans import Translator
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore

# Initialize Firebase
if not firebase_admin._apps:
    cred = credentials.Certificate("serviceAccountKey.json")
    firebase_admin.initialize_app(cred)
db = firestore.client()

# Initialize Google Translator
translator = Translator()

# Full Language Map: Full names to ISO codes
LANG_MAP = {
    "Arabic": "ar",
    "Chinese": "zh-cn",
    "Dutch": "nl",
    "English": "en",
    "French": "fr",
    "German": "de",
    "Hindi": "hi",
    "Italian": "it",
    "Japanese": "ja",
    "Korean": "ko",
    "Portuguese": "pt",
    "Russian": "ru",
    "Spanish": "es",
    "Turkish": "tr"
}

def get_language_code(name):
    return LANG_MAP.get(name, "en")

def detect_language(text):
    return detect(text)

def translate_text(text, src, dest):
    return translator.translate(text, src=src, dest=dest).text

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join([page.extract_text() or '' for page in reader.pages])

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def create_translated_pdf(text, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    font_path = os.path.join("assets", "fonts", "DejaVuSans.ttf")
    pdf.add_font("DejaVu", "", font_path, uni=True)
    pdf.set_font("DejaVu", size=12)

    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)

    pdf.output(output_path)


def create_translated_docx(text, output_path):
    doc = Document()
    for para in text.split('\n'):
        doc.add_paragraph(para)
    doc.save(output_path)

def save_txt(text, output_path):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def run(uid):
    st.title("📄 Document Translator")
    st.write("Upload a **PDF**, **DOCX**, **TXT**, **MD**, or **RTF** file to translate its contents.")

    uploaded_file = st.file_uploader("Choose a document", type=["pdf", "docx", "txt", "md", "rtf"])
    target_language = st.selectbox("Translate to", options=list(LANG_MAP.keys()))

    if uploaded_file and target_language:
        with st.spinner("🔍 Reading document..."):
            ext = uploaded_file.name.split('.')[-1].lower()
            if ext == "pdf":
                text = extract_text_from_pdf(uploaded_file)
            elif ext == "docx":
                text = extract_text_from_docx(uploaded_file)
            elif ext in ["txt", "md", "rtf"]:
                text = uploaded_file.read().decode("utf-8")
            else:
                st.error("❌ Unsupported file format.")
                return

        st.success("✅ Text extracted successfully!")
        st.text_area("Extracted Text Preview", text[:2000], height=200)

        if st.button("🌍 Translate Document"):
            with st.spinner("🌐 Translating..."):
                src_lang = detect_language(text)
                tgt_lang_code = get_language_code(target_language)
                translated = translate_text(text, src=src_lang, dest=tgt_lang_code)

                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as out_file:
                    out_path = out_file.name

                if ext == "pdf":
                    create_translated_pdf(translated, out_path)
                elif ext == "docx":
                    create_translated_docx(translated, out_path)
                else:
                    save_txt(translated, out_path)

                # Firestore save
                user_email = st.session_state.get("user_email", "anonymous")
                db.collection("translation_history").add({
                    "user": user_email,
                    "filename": uploaded_file.name,
                    "filetype": ext,
                    "source_lang": src_lang,
                    "target_lang": tgt_lang_code,
                    "timestamp": datetime.utcnow(),
                    "text_preview": text[:200]
                })

                st.success("✅ Document translated and saved!")
                with open(out_path, "rb") as f:
                    st.download_button(
                        f"📥 Download Translated {ext.upper()}",
                        f,
                        file_name=f"translated_{uploaded_file.name}"
                    )

    if st.button("⬅ Back to Dashboard"):
        st.session_state["current_page"] = "dashboard"
        st.rerun()

